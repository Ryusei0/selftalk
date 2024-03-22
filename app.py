# -*- coding: utf-8 -*-
from gevent import monkey; monkey.patch_all()
from datetime import datetime
import uuid
from flask import Flask, request, jsonify, send_file, abort, current_app, send_from_directory,  Response
from flask_cors import CORS
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import Qdrant
import openai
from openai import OpenAI
import os
import qdrant_client
import json
import azure.cognitiveservices.speech as speechsdk
from azure.cognitiveservices.speech import SpeechConfig, SpeechSynthesizer
from azure.cognitiveservices.speech.audio import AudioOutputConfig
from qdrant_client import QdrantClient, models
from threading import Timer
import boto3
from boto3.dynamodb.conditions import Key
import logging
from queue import Queue
from langchain_core.messages import HumanMessage, SystemMessage
from langchain_core.prompts.chat import (
    ChatPromptTemplate,
    HumanMessagePromptTemplate,
    SystemMessagePromptTemplate,
)
from langchain_openai import ChatOpenAI
from typing import Any, Dict, List, Optional, Union
from langchain.schema import AgentAction, AgentFinish, LLMResult
from langchain.callbacks.base import BaseCallbackHandler
from langchain.callbacks.manager import CallbackManager
import threading
import time 
from flask_socketio import SocketIO, emit
import tempfile
from docx import Document
import re

openai.api_key = os.environ['OPENAI_API_KEY']
qdrant_api_key = os.environ['QDRANT_KEY']
qdrant_url_key = os.environ['QDRANT_URL']

speech_key = os.environ['SPEECH_KEY']
service_region = os.environ['SERVICE_REGION_KEY']

s3 = boto3.client('s3')
BUCKET_NAME = 'testunity1.0'  # S3バケット名を設定
AWSreagion=os.environ['AWS_REGION_KEY']

app = Flask(__name__)
CORS(app)
# 全てのオリジンからの接続を許可する場合
socketio = SocketIO(app, async_mode='gevent', cors_allowed_origins="*")
message_queue = Queue()

person_id="ID-002"
company_id="001"

points=""

student=""

feature=""

# Qdrantクライアントの生成を関数外に移動
llm = OpenAI()

embeddings = OpenAIEmbeddings()

client = qdrant_client.QdrantClient(
    qdrant_url_key,
    api_key=qdrant_api_key,
)

global_contentss = []

 #以下は、リアルタイム返信
logging.basicConfig(level=logging.INFO)  # INFOレベル以上のログを出力

logger = logging.getLogger(__name__)

dynamodb = boto3.resource('dynamodb', region_name='ap-northeast-1')
table3 = dynamodb.Table('summary')
table2 = dynamodb.Table('conversations')
table = dynamodb.Table('maindatabase')
azure_speech_key = os.environ['SPEECH_KEY']
azure_service_region = os.environ['SERVICE_REGION_KEY']

message_queue = []  # クライアントのリスナーを追跡


db_path = 'conversation_database.db'

buffered_text = ""

def text_to_speech2(text,id):
    # 一時ファイルのパスを生成
    temp_file = tempfile.NamedTemporaryFile(delete=False)
    output_filename = f"talkaudio/text2output_{id}_{datetime.utcnow().strftime('%Y%m%d_%H%M%S%f')}_{uuid.uuid4().hex}.wav"
    file_path = temp_file.name

    speech_config = speechsdk.SpeechConfig(subscription=azure_speech_key, region=azure_service_region)
    audio_config = speechsdk.audio.AudioOutputConfig(filename=file_path)
    synthesizer = speechsdk.SpeechSynthesizer(speech_config=speech_config, audio_config=audio_config)

    ssml_string = f"""<speak version="1.0" xmlns="http://www.w3.org/2001/10/synthesis"
                    xmlns:mstts="https://www.w3.org/2001/mstts" xml:lang="ja-JP">
                    <voice name="ja-JP-NanamiNeural">
                        <mstts:express-as style="customerservice" styledegree="3">
                            {text}
                        </mstts:express-as>
                    </voice>
                  </speak>"""
    result = synthesizer.speak_ssml_async(ssml_string).get()

    # 音声データを一時ファイルに書き込み
    temp_file.close()

    # 一時ファイルをS3にアップロード
    with open(file_path, 'rb') as audio_file:
        s3.put_object(Bucket=BUCKET_NAME, Key=output_filename, Body=audio_file)
    
    # アップロードが完了したら一時ファイルを削除
    os.remove(file_path)
    
    # 正しいエンドポイントを反映したS3オブジェクトのURLを生成
    audio_url = f"https://s3.{AWSreagion}.amazonaws.com/{BUCKET_NAME}/{output_filename}"

    return result, audio_url

def openai_qa_univ(query, talk,talk2 ,id):
    print(f"Received query for QA: {query}")
    # CallbackManagerとMyCustomCallbackHandlerのインスタンス化時にidを渡す
    callback_handler = MyCustomCallbackHandler(lambda token, id=id: dummy_callback(token, id), id)
    callback_manager = CallbackManager([callback_handler])
    llm = ChatOpenAI(temperature=0, model_name="gpt-4-0125-preview", callback_manager=callback_manager, streaming=True)
    messages = [
    SystemMessage(
        content="You are a helpful assistant."
    ),
    HumanMessage(
        content=f"""あなたは聞き上手な、キャリア相談のプロフェッショナルです。

「場面」における、あなたの「役割」を、「ルール１」、「ルール２」、「ルール３」、「ルール４」、「ルール５」、「ルール６」、「ルール７」、「ルール８」、「ルール９」に従い遂行し、「ゴール」を達成してください。

#条件

「場面」：キャリアに悩む学生が、あなたのところへ相談に来ている。

「役割」：メンターとして学生と日常会話をし、学生の話を聞く。学生はあなたしか頼れる人物がいないので、学生に質問を投げかけ、熱心に学生のことを理解する、最高の頼れるメンターとしての役割を全うする。

「ゴール」：会話履歴の一番最後の行の、「あなた：」のセリフを埋める。（このセリフは、映画の台本として使用する。）

「ルール１」：会話は、日常会話なので、セリフは多くても数行に抑える。

「ルール２」：この後も会話は続くので、会話は絶対に完結させないように自然なセリフを考える。

「ルール３」：あなたは、学生が少しでも「学生の質問の真意」を把握できるように、深掘りを重視し、多様な話を質問をする。

「ルール４」：以下の＃会話履歴は、あなたと学生の過去の会話なので、最大限会話の内容を考慮し、同じ内容を聞く質問をしない。
             
「ルール４」：回答は「あなた：」より後ろのセリフのみ出力する。

「ルール５」：あなたは以下の#性格を持っている。

「ルール６」：以下の#口調のように、セリフは若い女性の喋り口調に最大限似せる。

「ルール７」：「お前」という言葉は決して使わない。

「ルール８」：学生の名前は「りゅうせい」なので、名前を呼ぶときに使用する。

「ルール９」：会話の最後には、学生への問いかけを入れる。

#口調
これかなり面白いよ！（フレンドリーな話し方）

#性格
・学生のことが好きで、学生の力になりたい
・学生のことを知り、興味を引き出したい
・学生の興味の深掘りがしたい
・学生の会話をリードしてあげたい

#学生の情報

{talk2}

#会話履歴

{talk}

学生:{query}

あなた："""
    ),
]
    response = llm.invoke(messages)
    print(f"LLM response: {response}")
    return response

def openai_summary(query, talk,talk2,id):
    print(f"Received query for QA: {query}")
    # CallbackManagerとMyCustomCallbackHandlerのインスタンス化時にidを渡す
    callback_handler = MyCustomCallbackHandler(lambda token, id=id: dummy_callback(token, id), id)
    callback_manager = CallbackManager([callback_handler])
    llm = ChatOpenAI(temperature=0, model_name="gpt-4-0125-preview", callback_manager=callback_manager, streaming=True)
    messages = [
    SystemMessage(
        content="You are a helpful assistant."
    ),
    HumanMessage(
        content=f"""あなたは聞き上手な、キャリア相談のプロフェッショナルです。

「場面」における、あなたの「役割」を、「ルール１」、「ルール２」、「ルール３」、「ルール４」、「ルール５」、「ルール６」、「ルール７」、「ルール８」、「ルール９」に従い遂行し、「ゴール」を達成してください。

#条件

「場面」：キャリアに悩む学生が、あなたのところへ相談に来ている。（学生は以前もあなたのもとを訪れており、以下の情報は聞き出している。）

「役割」：メンターとして学生と日常会話をし、学生の話を聞く。学生はあなたしか頼れる人物がいないので、以下の「知りたい項目」が知れるように、学生に質問を投げかけ、熱心に学生のことを理解する、最高の頼れるメンターとしての役割を全うする。ただし、すでに学生との会話で以下の情報は聞き出しているので、「知りたい項目」の中で足りていない情報を聞き出せるように質問する。

「ゴール」：会話履歴の一番最後の行の、「あなた：」のセリフを埋める。（このセリフは、映画の台本として使用する。）

「ルール１」：会話は、日常会話なので、セリフは多くても数行に抑える。（100文字程度）

「ルール２」：この後も会話は続くので、会話は絶対に完結させないように自然なセリフを考える。

「ルール３」：あなたは、学生が少しでも学生の特徴を把握できるように、深掘りを重視し、多様な話を質問をする。

「ルール４」：以下の＃会話履歴は、あなたと学生の過去の会話なので、最大限会話の内容を考慮し、同じ内容を聞く質問をしない。
             
「ルール４」：回答は「あなた：」より後ろのセリフのみ出力する。

「ルール５」：あなたは以下の#性格を持っている。

「ルール６」：以下の#口調のように、セリフは若い女性の喋り口調に最大限似せる。

「ルール７」：「お前」という言葉は決して使わない。

「ルール８」：学生の名前は「りゅうせい」なので、名前を呼ぶときに使用する。

「ルール９」：会話の最後には、学生への問いかけを入れる。

#口調
これかなり面白いよ！（フレンドリーな話し方）

#性格
・学生のことが好きで、学生の力になりたい
・学生のことを知り、興味を引き出したい
・学生の興味の深掘りがしたい
・学生の会話をリードしてあげたい

#あなたが知りたい項目　

1. **基本情報**
   - 年齢
   - 教育歴
   - 職務経験
   - 現在の職業状況（就職中、求職中等）
   - 転職回数とその理由

2. **能力とスキル**
   - 専門的なスキルや知識
   - 獲得した資格や認定
   - コミュニケーション能力
   - リーダーシップやチームワークの能力
   - 問題解決能力
   - 学習意欲と自己啓発への姿勢

3. **興味関心**
   - 職業上の興味や好奇心がある分野
   - ホビー、趣味、活動への関心
   - 仕事における好みや不得意なこと

4. **価値観と優先順位**
   - キャリアにおける価値観（例：報酬、社会貢献、ワークライフバランス）
   - 将来的な目標と期待（短期、中期、長期）
   - 優先順位（家族、余暇、キャリア進展）

5. **生活状況と制約**
   - 地理的な制約（転居の可否）
   - 家族状況（配偶者の職業、子育て等）
   - 経済的な制約（最低限必要な収入）
   - 健康状態や身体的な制約

6. **キャリアの選好**
   - 望む職種や業種
   - 希望する職場の文化や環境
   - 避けたい業務や職場環境
   - 働き方の好み（フルタイム、パートタイム、リモートワーク）

#以前に聞き出した学生の情報

{talk2}

#会話履歴

{talk}

学生:{query}

あなた："""
    ),
]
    response = llm.invoke(messages)
    print(f"LLM response: {response}")
    return response


def openai_restart(talk,id):
    print(f"Received restart")
    llm = ChatOpenAI(temperature=0, model_name="gpt-4-0125-preview")
    messages = [
    SystemMessage(
        content="You are a helpful assistant."
    ),
    HumanMessage(
        content=f"""あなたはキャリア相談のプロフェッショナルです。

以下のゴールを達成してください。

#ゴール

以下の会話履歴から、学生の情報を正確に抜き出し、以下の整理項目に従って、箇条書きでまとめる。

#条件

・会話履歴から確実に読み取れることをまとめること。

・出力は、以下の出力例に合わせること。

＃整理項目

1. **基本情報**（年齢、教育歴、職歴など）

2. **能力とスキル**

3. **興味関心**

4. **価値観と優先順位**（キャリアにおける価値観や、将来的な目標と期待、余暇かキャリア進展のどちらを優先するかなど。）

5. **生活状況と制約**

6. **キャリアの選好**

#出力例

1. **基本情報**
・
・
、、、

2. **能力とスキル**
・
・
、、、

3. **興味関心**
・
・
、、、

4. **価値観と優先順位**（キャリアにおける価値観や、将来的な目標と期待、余暇かキャリア進展のどちらを優先するかなど。）
・
・
、、、

5. **生活状況と制約**
・
・
、、、

6. **キャリアの選好**
・
・
、、、

#会話履歴

{talk}
"""
    ),
]
    response = llm.invoke(messages)
    print(f"LLM response: {response}")
    return response

def openai_company(talk,id):
    print(f"Received restart")
    llm = ChatOpenAI(temperature=0, model_name="gpt-4-0125-preview")
    messages = [
    SystemMessage(
        content="You are a helpful assistant."
    ),
    HumanMessage(
        content=f"""あなたはキャリア設計のプロフェッショナルです。

以下のゴールを達成してください。

#ゴール

以下の学生の特徴を基に、学生におすすめの企業を具体的にリストアップすること

#条件

・出力は、以下の出力例に従う。

・おすすめ企業は、「大企業」、「中小企業、ベンチャー企業」毎にそれぞれ3社提案する。

・おすすめ分野を３つ提案し、それぞれの分野で、「大企業」と「中小企業、ベンチャー企業」を3社ずつ、計18社具体的に提案する。

#出力例

おすすめ1:コンサルティング領域

「大企業」
・
・
・

「中小企業、ベンチャー企業」
・
・
・

おすすめ2:
、、、

おすすめ3:
、、、

#学生の特徴

{talk}

"""
    ),
]
    response = llm.invoke(messages)
    print(f"LLM response: {response}")
    return response

@socketio.on('textuniv')
def handle_messageuniv(data):
    query = data.get('text')
    id = data.get('id')  # ユーザーIDの取得
    user_id = "001"
    if not query:
        return "No text provided", 400
    try:
       # DynamoDBから最新の会話履歴を取得
        talk = get_latest_conversation(user_id) if get_latest_conversation(user_id) else ""
        talk2 = get_latest_summary(user_id) if get_latest_summary(user_id) else ""
        ai_message = openai_qa_univ(query, talk,talk2 ,id)
        # AIMessage オブジェクトの content 属性からテキスト内容を抽出
        response_text = ai_message.content if ai_message and hasattr(ai_message, "content") else "応答を取得できませんでした。"
        fix= f"{talk}\n学生: {query}\nあなた: {response_text}"
        # DynamoDBへの保存
        save_conversation(user_id, query, response_text, fix)
        # クライアントにレスポンステキストを送信
        socketio.emit(f'response_{id}', {'response': response_text})
        # 応答テキストをクライアントに返す
        return jsonify({'response': response_text}), 200
    except Exception as e:
        logger.error(f"Error processing request: {e}")
        return jsonify({"error": "Internal server error"}), 500
    
@socketio.on('specialEndpoint')
def summary_handle_messageuniv(data):
    query = data.get('text')
    id = data.get('id')  # ユーザーIDの取得
    user_id = "001"
    if not query:
        return "No text provided", 400
    try:
       # DynamoDBから最新の会話履歴を取得
        talk = get_latest_conversation(user_id) if get_latest_conversation(user_id) else ""
        talk2 = get_latest_summary(user_id) if get_latest_summary(user_id) else ""
        ai_message = openai_summary(query, talk,talk2 ,id)
        # AIMessage オブジェクトの content 属性からテキスト内容を抽出
        response_text = ai_message.content if ai_message and hasattr(ai_message, "content") else "応答を取得できませんでした。"
        fix= f"{talk}\n学生: {query}\nあなた: {response_text}"
        # DynamoDBへの保存
        save_conversation(user_id, query, response_text, fix)
        # クライアントにレスポンステキストを送信
        socketio.emit(f'response_{id}', {'response': response_text})
        # 応答テキストをクライアントに返す
        return jsonify({'response': response_text}), 200
    except Exception as e:
        logger.error(f"Error processing request: {e}")
        return jsonify({"error": "Internal server error"}), 500

def get_latest_conversation(user_id):
    try:
        response = table2.query(
            KeyConditionExpression=Key('conversationId').eq(user_id),
            ScanIndexForward=False,  # 最新の項目を最初にする
            Limit=1
        )
        if response['Items']:
            # 'conversation'キーの値のみを返す
            return response['Items'][0]['conversation']
        else:
            return None
    except Exception as e:
        logger.error(f"Failed to get latest conversation for user {user_id}: {str(e)}")
        return None

def get_latest_summary(user_id):
    try:
        response = table3.query(
            KeyConditionExpression=Key('conversationId').eq(user_id),
            ScanIndexForward=False,  # 最新の項目を最初にする
            Limit=1
        )
        if response['Items']:
            # 'conversation'キーの値のみを返す
            return response['Items'][0]['summary']
        else:
            return None
    except Exception as e:
        logger.error(f"Failed to get latest conversation for user {user_id}: {str(e)}")
        return None

@socketio.on('restart')
def restart_messageuniv(data):
    id = data.get('id')  # ユーザーIDの取得
    user_id = "001"
    if not id:
        return "No text provided", 400
    try:
       # DynamoDBから最新の会話履歴を取得
        talk = get_latest_conversation(user_id) if get_latest_conversation(user_id) else ""
        talk2 = get_latest_summary(user_id) if get_latest_summary(user_id) else ""
        ai_message = openai_restart(talk, id)
        # AIMessage オブジェクトの content 属性からテキスト内容を抽出
        response_text = ai_message.content if ai_message and hasattr(ai_message, "content") else "応答を取得できませんでした。"
        # DynamoDBへの保存　ここを整理（ユーザーの特徴）その後保存して、削除
        save_conversation_delete(user_id,talk,talk2,response_text)
        socketio.emit(f'response_{id}', {'response': response_text})
        # 応答テキストをクライアントに返す
        return jsonify({'response': response_text}), 200
    except Exception as e:
        logger.error(f"Error processing request: {e}")
        return jsonify({"error": "Internal server error"}), 500
    
@socketio.on('company')
def company_messageuniv(data):
    id = data.get('id')  # ユーザーIDの取得
    user_id = "001"
    if not id:
        return "No text provided", 400
    try:
       # DynamoDBから最新の会話履歴を取得
        talk = get_latest_conversation(user_id) if get_latest_conversation(user_id) else ""
        talk2 = get_latest_summary(user_id) if get_latest_summary(user_id) else ""
        ai_message = openai_restart(talk, id)
        # AIMessage オブジェクトの content 属性からテキスト内容を抽出
        response_text = ai_message.content if ai_message and hasattr(ai_message, "content") else "応答を取得できませんでした。"
        # DynamoDBへの保存　ここを整理（ユーザーの特徴）その後保存して、削除
        comsummary=save_conversation_delete(user_id,talk,talk2,response_text)
        ai_message2 = openai_company(comsummary, id)
        response_text2 = ai_message2.content if ai_message2 and hasattr(ai_message2, "content") else "応答を取得できませんでした。"
        saveward(user_id, response_text2)
        socketio.emit(f'response_{id}', {'response': response_text2})
        # 応答テキストをクライアントに返す
        return jsonify({'response': response_text}), 200
    except Exception as e:
        logger.error(f"Error processing request: {e}")
        return jsonify({"error": "Internal server error"}), 500
    
def saveward(user_id, conversation):
     # 新しいワード文書を作成
     doc = Document()
     doc.add_heading('Conversation Summary', 0)
     doc.add_paragraph('User ID: ' + user_id)
     doc.add_page_break()
     doc.add_heading('Full Conversation', level=1)
     doc.add_paragraph(conversation)

    # ファイル名を設定（ユーザーIDと現在のタイムスタンプを使用）
     filename = f'conversation_{user_id}_{int(time.time())}.docx'
     local_path = f'/tmp/{filename}'

    # 文書をローカルに保存
     doc.save(local_path)
 
    # S3にアップロード
     s3.upload_file(local_path, BUCKET_NAME, filename)

    # ローカルのファイルを削除
     os.remove(local_path)
    
#会話を保存、要約、削除 
def save_conversation_delete(user_id, conversation,presummary,summary):
    timestamp = datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')
    comsummary=extract_and_combine(presummary,summary)
    try:
     saveward(user_id, conversation)

     response = table3.put_item(
            Item={
                'conversationId': user_id,  # ユーザーIDを直接会話IDとして使用
                'timestamp': timestamp,  # このタイムスタンプを使って最新のレコードを識別
                'summary': comsummary,
            }
        )
     logger.info(f"Conversation for user {user_id} saved successfully.")
     # 'conversations' テーブルからアイテムを取得
     items = scan_dynamodb_table('conversations')

# 取得したアイテムを削除
     delete_items_from_dynamodb_table('conversations', items)
     return comsummary
    except Exception as e:
        logger.error(f"Failed to save conversation for user {user_id}: {str(e)}")

def save_conversation(user_id, query, final_response, conversation):
    timestamp = datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')
    
    try:
        response = table2.put_item(
            Item={
                'conversationId': user_id,  # ユーザーIDを直接会話IDとして使用
                'timestamp': timestamp,  # このタイムスタンプを使って最新のレコードを識別
                'query': query,
                'final_response': final_response,
                'conversation': conversation
            }
        )
        logger.info(f"Conversation for user {user_id} saved successfully.")
    except Exception as e:
        logger.error(f"Failed to save conversation for user {user_id}: {str(e)}")

def scan_dynamodb_table(table_name):
    dynamodb = boto3.resource('dynamodb', region_name='ap-northeast-1')
    table = dynamodb.Table(table_name)

    # テーブルをスキャンして全アイテムを取得
    response = table.scan()

    # アイテムのリストを取得
    items = response['Items']

    return items

def delete_items_from_dynamodb_table(table_name, items):
    dynamodb = boto3.resource('dynamodb', region_name='ap-northeast-1')
    table = dynamodb.Table(table_name)

    # アイテムごとに削除処理
    for item in items:
        # conversationIdとtimestampでアイテムを特定して削除
        response = table.delete_item(
            Key={
                'conversationId': item['conversationId'],
                'timestamp': item['timestamp']  # 新しいテーブル構造に合わせてフィールド名を調整
            }
        )
        print(response)


def dummy_callback(token, id):
    global buffered_text
    buffered_text += token
    print(f'callback>> {token}', end='')

    if token.endswith("。") or token.endswith("?") or token.endswith("!") or token.endswith("？") or token.endswith("！"):
        print("\nToken ends with a period. Generating audio...")
        text = buffered_text
        buffered_text = ""  # バッファをクリア
        result, output_filename = text_to_speech2(text, id)
        print("Message successfully added to queue.")
        audio_url =  output_filename
        # 変更点: ユーザー固有のイベントに対してemit
        socketio.emit(f'audio_url_{id}', {'url': audio_url})
        print(f'audio_url_{id}:Message send')

class MyCustomCallbackHandler(BaseCallbackHandler):
    def __init__(self, callback, id):  # idを受け取るように変更
        self.callback = callback
        self.id = id  # idをインスタンス変数として保存
    def on_llm_start(
        self, serialized: Dict[str, Any], prompts: List[str], **kwargs: Any
    ) -> Any:
        """Run when LLM starts running."""
    def on_llm_new_token(self, token: str, **kwargs: Any) -> Any:
        """Run on new LLM token. Only available when streaming is enabled."""
        if self.callback is not None:
            self.callback(token, self.id)  # idも渡してcallbackを呼び出す
    def on_llm_end(self, response: LLMResult, **kwargs: Any) -> Any:
        """Run when LLM ends running."""
    def on_llm_error(
        self, error: Union[Exception, KeyboardInterrupt], **kwargs: Any
    ) -> Any:
        """Run when LLM errors."""
    def on_chain_start(
        self, serialized: Dict[str, Any], inputs: Dict[str, Any], **kwargs: Any
    ) -> Any:
        """Run when chain starts running."""
    def on_chain_end(self, outputs: Dict[str, Any], **kwargs: Any) -> Any:
        """Run when chain ends running."""
    def on_chain_error(
        self, error: Union[Exception, KeyboardInterrupt], **kwargs: Any
    ) -> Any:
        """Run when chain errors."""
    def on_tool_start(
        self, serialized: Dict[str, Any], input_str: str, **kwargs: Any
    ) -> Any:
        """Run when tool starts running."""
    def on_tool_end(self, output: str, **kwargs: Any) -> Any:
        """Run when tool ends running."""
    def on_tool_error(
        self, error: Union[Exception, KeyboardInterrupt], **kwargs: Any
    ) -> Any:
        """Run when tool errors."""
    def on_text(self, text: str, **kwargs: Any) -> Any:
        """Run on arbitrary text."""
    def on_agent_action(self, action: AgentAction, **kwargs: Any) -> Any:
        """Run on agent action."""
    def on_agent_finish(self, finish: AgentFinish, **kwargs: Any) -> Any:
        """Run on agent end."""


def extract_and_combine(text1, text2):
    combined_info = {}
    # 項目を抽出する正規表現パターンを調整し、数字と「.」に基づいて分割するようにします。
    pattern = r"(\d+)\. \*\*(.*?)\*\*(.+?)(?=\n\d+\. |\n$)"

    # 両方のテキストから情報を抽出
    items1 = re.findall(pattern, text1, re.DOTALL)
    items2 = re.findall(pattern, text2, re.DOTALL)

    # 抽出した情報を辞書に統合する
    for item in items1 + items2:
        num, title, content = item
        # コンテンツの前後の空白を削除
        content = re.sub(r'^\s+|\s+$', '', content, flags=re.MULTILINE)
        if title not in combined_info:
            combined_info[title] = content
        else:
            # 重複を避けつつ新しい情報を追加
            existing_content = combined_info[title]
            new_contents = [line.strip() for line in content.split('\n') if line.strip() and line.strip() not in existing_content]
            if new_contents:
                combined_info[title] += "\n   - " + "\n   - ".join(new_contents)

    return combined_info

@app.route('/')
def index():
    return 'Server is running!'

if __name__ == '__main__':
    socketio.run(app, host='0.0.0.0', port=8000, debug=True)
